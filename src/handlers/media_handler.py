"""Media handler for image and shape operations.

This module provides functionality for inserting and managing
images and other media in DOCX documents.
"""

import io
from pathlib import Path

from docx import Document
from docx.shared import Inches
from PIL import Image

from src.core.constants import MAX_IMAGE_DIMENSION, SUPPORTED_IMAGE_FORMATS
from src.core.exceptions import UnsupportedFormatError, ValidationError
from src.models.dto import ImageDTO


class MediaHandler:
    """Handler for media operations.

    This class provides methods for inserting and managing
    images and other media in DOCX documents.
    """

    def __init__(self, document: Document | None = None) -> None:
        """Initialize the media handler.

        Args:
            document: The Document instance to work with (optional).
        """
        self._document = document

    @property
    def document(self) -> Document:
        """Get the document instance."""
        if self._document is None:
            raise ValueError("No document loaded")
        return self._document

    def set_document(self, document: Document) -> None:
        """Set the document instance.

        Args:
            document: The Document instance to work with.
        """
        self._document = document

    def insert_image(
        self,
        image_path: str | Path,
        paragraph_index: int | None = None,
        width: float | None = None,
        height: float | None = None,
        alt_text: str | None = None,
    ) -> int:
        """Insert an image into the document.

        Args:
            image_path: Path to the image file.
            paragraph_index: Optional paragraph index for insertion.
            width: Optional width in inches.
            height: Optional height in inches.
            alt_text: Optional alternative text.

        Returns:
            Index of the inserted inline shape.

        Raises:
            UnsupportedFormatError: If the image format is not supported.
            ValidationError: If the image is invalid.
        """
        path = Path(image_path)
        if not path.exists():
            raise ValidationError(f"Image file not found: {image_path}")

        ext = path.suffix.lower()
        if ext not in SUPPORTED_IMAGE_FORMATS:
            raise UnsupportedFormatError(
                f"Unsupported image format: {ext}",
                format_=ext,
                supported_formats=SUPPORTED_IMAGE_FORMATS,
            )

        # Validate image dimensions
        with Image.open(path) as img:
            if img.width > MAX_IMAGE_DIMENSION or img.height > MAX_IMAGE_DIMENSION:
                raise ValidationError(
                    f"Image dimensions exceed maximum ({MAX_IMAGE_DIMENSION}px)"
                )

        # Prepare size arguments
        size_kwargs = {}
        if width is not None:
            size_kwargs["width"] = Inches(width)
        if height is not None:
            size_kwargs["height"] = Inches(height)

        # Insert image
        if paragraph_index is not None:
            if paragraph_index < 0 or paragraph_index >= len(self._document.paragraphs):
                raise ValidationError(f"Paragraph index {paragraph_index} out of range")
            para = self._document.paragraphs[paragraph_index]
            run = para.add_run()
            run.add_picture(str(path), **size_kwargs)
        else:
            para = self._document.add_paragraph()
            run = para.add_run()
            run.add_picture(str(path), **size_kwargs)

        return len(self._document.inline_shapes) - 1

    def insert_image_from_bytes(
        self,
        image_data: bytes,
        filename: str,
        paragraph_index: int | None = None,
        width: float | None = None,
        height: float | None = None,
    ) -> int:
        """Insert an image from bytes into the document.

        Args:
            image_data: Image data as bytes.
            filename: Original filename for extension detection.
            paragraph_index: Optional paragraph index for insertion.
            width: Optional width in inches.
            height: Optional height in inches.

        Returns:
            Index of the inserted inline shape.

        Raises:
            ValidationError: If the image is invalid.
        """
        ext = Path(filename).suffix.lower()
        if ext not in SUPPORTED_IMAGE_FORMATS:
            raise UnsupportedFormatError(
                f"Unsupported image format: {ext}",
                format_=ext,
                supported_formats=SUPPORTED_IMAGE_FORMATS,
            )

        # Validate image
        try:
            with Image.open(io.BytesIO(image_data)) as img:
                if img.width > MAX_IMAGE_DIMENSION or img.height > MAX_IMAGE_DIMENSION:
                    raise ValidationError(
                        f"Image dimensions exceed maximum ({MAX_IMAGE_DIMENSION}px)"
                    )
        except Exception as e:
            raise ValidationError(f"Invalid image data: {e}")

        # Prepare size arguments
        size_kwargs = {}
        if width is not None:
            size_kwargs["width"] = Inches(width)
        if height is not None:
            size_kwargs["height"] = Inches(height)

        # Insert image
        image_stream = io.BytesIO(image_data)

        if paragraph_index is not None:
            if paragraph_index < 0 or paragraph_index >= len(self._document.paragraphs):
                raise ValidationError(f"Paragraph index {paragraph_index} out of range")
            para = self._document.paragraphs[paragraph_index]
            run = para.add_run()
            run.add_picture(image_stream, **size_kwargs)
        else:
            para = self._document.add_paragraph()
            run = para.add_run()
            run.add_picture(image_stream, **size_kwargs)

        return len(self._document.inline_shapes) - 1

    def get_image_count(self) -> int:
        """Get the number of inline images in the document.

        Returns:
            Number of inline shapes (images).
        """
        return len(self._document.inline_shapes)

    def get_image_info(self, index: int) -> ImageDTO:
        """Get information about an inline image.

        Args:
            index: Index of the inline shape.

        Returns:
            Image DTO with information.

        Raises:
            ValidationError: If the index is out of range.
        """
        if index < 0 or index >= len(self._document.inline_shapes):
            raise ValidationError(f"Image index {index} out of range")

        shape = self._document.inline_shapes[index]

        # Get dimensions in inches
        width = shape.width.inches if shape.width else None
        height = shape.height.inches if shape.height else None

        return ImageDTO(
            index=index,
            paragraph_index=0,  # Would need more complex logic to determine
            filename="",
            width=width,
            height=height,
        )

    def resize_image(
        self,
        index: int,
        width: float | None = None,
        height: float | None = None,
    ) -> None:
        """Resize an inline image.

        Args:
            index: Index of the inline shape.
            width: New width in inches.
            height: New height in inches.

        Raises:
            ValidationError: If the index is out of range.
        """
        if index < 0 or index >= len(self._document.inline_shapes):
            raise ValidationError(f"Image index {index} out of range")

        shape = self._document.inline_shapes[index]

        if width is not None:
            shape.width = Inches(width)
        if height is not None:
            shape.height = Inches(height)

    def delete_image(self, index: int) -> None:
        """Delete an inline image.

        Args:
            index: Index of the inline shape to delete.

        Raises:
            ValidationError: If the index is out of range.
        """
        if index < 0 or index >= len(self._document.inline_shapes):
            raise ValidationError(f"Image index {index} out of range")

        shape = self._document.inline_shapes[index]
        # Remove the inline shape element from its parent
        inline = shape._inline
        inline.getparent().remove(inline)

    def compress_image(
        self,
        image_data: bytes,
        quality: int = 85,
        max_width: int | None = None,
        max_height: int | None = None,
    ) -> bytes:
        """Compress an image.

        Args:
            image_data: Original image data.
            quality: JPEG quality (1-100).
            max_width: Maximum width in pixels.
            max_height: Maximum height in pixels.

        Returns:
            Compressed image data as bytes.
        """
        with Image.open(io.BytesIO(image_data)) as img:
            # Resize if necessary
            if max_width or max_height:
                ratio = 1.0
                if max_width and img.width > max_width:
                    ratio = min(ratio, max_width / img.width)
                if max_height and img.height > max_height:
                    ratio = min(ratio, max_height / img.height)

                if ratio < 1.0:
                    new_width = int(img.width * ratio)
                    new_height = int(img.height * ratio)
                    img = img.resize((new_width, new_height), Image.Resampling.LANCZOS)

            # Convert and compress
            output = io.BytesIO()
            if img.mode in ("RGBA", "P"):
                img = img.convert("RGB")
            img.save(output, format="JPEG", quality=quality, optimize=True)
            output.seek(0)
            return output.read()

    def add_text_box(
        self,
        text: str,
        width: float,
        height: float,
        paragraph_index: int | None = None,
    ) -> None:
        """Add a text box to the document.

        Args:
            text: Text content for the box.
            width: Width in inches.
            height: Height in inches.
            paragraph_index: Optional paragraph index for insertion.

        Note:
            This is a placeholder - python-docx has limited text box support.
        """
        # python-docx doesn't have direct text box support
        # We'll add a paragraph with the text instead
        if paragraph_index is not None:
            if paragraph_index < 0 or paragraph_index >= len(self._document.paragraphs):
                raise ValidationError(f"Paragraph index {paragraph_index} out of range")
        self._document.add_paragraph(text)
