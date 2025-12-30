"""TOC handler for table of contents and navigation operations.

This module provides functionality for managing table of contents,
bookmarks, and hyperlinks in DOCX documents.
"""

from typing import Any

from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn

from src.core.exceptions import ValidationError
from src.models.dto import BookmarkDTO, HyperlinkDTO


class TocHandler:
    """Handler for TOC and navigation operations.

    This class provides methods for managing table of contents,
    bookmarks, and hyperlinks in DOCX documents.
    """

    def __init__(self, document: Document | None = None) -> None:
        """Initialize the TOC handler.

        Args:
            document: The Document instance to work with (optional).
        """
        self._document = document

    @property
    def document(self) -> Document:
        """Get the document instance."""
        if self._document is None:
            raise ValueError("No document loaded")
        return self._document

    def set_document(self, document: Document) -> None:
        """Set the document instance.

        Args:
            document: The Document instance to work with.
        """
        self._document = document

    def add_table_of_contents(
        self,
        title: str = "Table of Contents",
        max_level: int = 3,
        paragraph_index: int | None = None,
    ) -> int:
        """Add a table of contents to the document.

        Args:
            title: TOC title.
            max_level: Maximum heading level to include.
            paragraph_index: Optional position to insert TOC.

        Returns:
            Index of the TOC paragraph.

        Note:
            The TOC will need to be updated in Word to populate.
        """
        # Add title
        toc_title = self._document.add_paragraph(title)
        toc_title.style = "Heading 1"

        # Add TOC field
        para = self._document.add_paragraph()
        run = para.add_run()

        # Create TOC field code
        fld_char_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        instr_text = parse_xml(
            f'<w:instrText {nsdecls("w")} xml:space="preserve"> TOC \\o "1-{max_level}" \\h \\z \\u </w:instrText>'
        )
        fld_char_separate = parse_xml(
            f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>'
        )
        fld_char_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')

        run._r.append(fld_char_begin)
        run._r.append(instr_text)
        run._r.append(fld_char_separate)
        run._r.append(fld_char_end)

        return len(self._document.paragraphs) - 1

    def update_toc(self) -> None:
        """Mark the TOC for update.

        Note:
            The actual update happens when the document is opened in Word.
        """
        # TOC update requires Word application to process
        # We can only set flags to indicate update is needed
        pass

    def add_bookmark(
        self,
        name: str,
        paragraph_index: int,
    ) -> BookmarkDTO:
        """Add a bookmark to the document.

        Args:
            name: Bookmark name (must be unique).
            paragraph_index: Index of the paragraph to bookmark.

        Returns:
            Created bookmark DTO.

        Raises:
            ValidationError: If the index is out of range.
        """
        if paragraph_index < 0 or paragraph_index >= len(self._document.paragraphs):
            raise ValidationError(f"Paragraph index {paragraph_index} out of range")

        para = self._document.paragraphs[paragraph_index]
        p = para._p

        # Generate unique ID
        import random

        bookmark_id = str(random.randint(1000000, 9999999))

        # Create bookmark start element
        bookmark_start = parse_xml(
            f'<w:bookmarkStart {nsdecls("w")} w:id="{bookmark_id}" w:name="{name}"/>'
        )
        # Create bookmark end element
        bookmark_end = parse_xml(
            f'<w:bookmarkEnd {nsdecls("w")} w:id="{bookmark_id}"/>'
        )

        # Insert at the beginning of the paragraph
        p.insert(0, bookmark_start)
        p.append(bookmark_end)

        return BookmarkDTO(
            name=name,
            paragraph_index=paragraph_index,
        )

    def get_bookmarks(self) -> list[BookmarkDTO]:
        """Get all bookmarks in the document.

        Returns:
            List of bookmark DTOs.
        """
        bookmarks = []
        bookmark_names: dict[str, int] = {}

        for i, para in enumerate(self._document.paragraphs):
            # Find bookmark starts in this paragraph
            for elem in para._p.iterchildren():
                if elem.tag == qn("w:bookmarkStart"):
                    name = elem.get(qn("w:name"))
                    if name and name not in bookmark_names:
                        bookmark_names[name] = i

        for name, index in bookmark_names.items():
            bookmarks.append(BookmarkDTO(name=name, paragraph_index=index))

        return bookmarks

    def delete_bookmark(self, name: str) -> None:
        """Delete a bookmark by name.

        Args:
            name: Bookmark name to delete.

        Raises:
            ValidationError: If the bookmark is not found.
        """
        found = False
        bookmark_id = None

        # Find and remove bookmark elements
        for para in self._document.paragraphs:
            elements_to_remove = []
            for elem in para._p.iterchildren():
                if elem.tag == qn("w:bookmarkStart"):
                    if elem.get(qn("w:name")) == name:
                        bookmark_id = elem.get(qn("w:id"))
                        elements_to_remove.append(elem)
                        found = True
                elif elem.tag == qn("w:bookmarkEnd"):
                    if elem.get(qn("w:id")) == bookmark_id:
                        elements_to_remove.append(elem)

            for elem in elements_to_remove:
                para._p.remove(elem)

        if not found:
            raise ValidationError(f"Bookmark not found: {name}")

    def add_hyperlink(
        self,
        text: str,
        url: str,
        paragraph_index: int,
        offset: int | None = None,
    ) -> HyperlinkDTO:
        """Add a hyperlink to the document.

        Args:
            text: Link text.
            url: Link URL.
            paragraph_index: Index of the paragraph.
            offset: Optional character offset for insertion.

        Returns:
            Created hyperlink DTO.

        Raises:
            ValidationError: If the index is out of range.
        """
        if paragraph_index < 0 or paragraph_index >= len(self._document.paragraphs):
            raise ValidationError(f"Paragraph index {paragraph_index} out of range")

        para = self._document.paragraphs[paragraph_index]

        # Add hyperlink relationship
        part = self._document.part
        r_id = part.relate_to(
            url,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True,
        )

        # Create hyperlink element
        hyperlink = parse_xml(
            f'<w:hyperlink {nsdecls("w")} r:id="{r_id}" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
            f"<w:r>"
            f'<w:rPr><w:rStyle w:val="Hyperlink"/></w:rPr>'
            f"<w:t>{text}</w:t>"
            f"</w:r>"
            f"</w:hyperlink>"
        )

        para._p.append(hyperlink)

        return HyperlinkDTO(
            text=text,
            url=url,
            paragraph_index=paragraph_index,
        )

    def add_internal_link(
        self,
        text: str,
        bookmark_name: str,
        paragraph_index: int,
    ) -> HyperlinkDTO:
        """Add an internal link to a bookmark.

        Args:
            text: Link text.
            bookmark_name: Name of the target bookmark.
            paragraph_index: Index of the paragraph.

        Returns:
            Created hyperlink DTO.

        Raises:
            ValidationError: If the index is out of range.
        """
        if paragraph_index < 0 or paragraph_index >= len(self._document.paragraphs):
            raise ValidationError(f"Paragraph index {paragraph_index} out of range")

        para = self._document.paragraphs[paragraph_index]

        # Create internal hyperlink element
        hyperlink = parse_xml(
            f'<w:hyperlink {nsdecls("w")} w:anchor="{bookmark_name}">'
            f"<w:r>"
            f'<w:rPr><w:rStyle w:val="Hyperlink"/></w:rPr>'
            f"<w:t>{text}</w:t>"
            f"</w:r>"
            f"</w:hyperlink>"
        )

        para._p.append(hyperlink)

        return HyperlinkDTO(
            text=text,
            url=f"#{bookmark_name}",
            paragraph_index=paragraph_index,
        )

    def get_hyperlinks(self) -> list[HyperlinkDTO]:
        """Get all hyperlinks in the document.

        Returns:
            List of hyperlink DTOs.
        """
        hyperlinks = []

        for i, para in enumerate(self._document.paragraphs):
            for elem in para._p.iterchildren():
                if elem.tag == qn("w:hyperlink"):
                    # Get text from the hyperlink
                    text_parts = []
                    for t in elem.iter(qn("w:t")):
                        if t.text:
                            text_parts.append(t.text)
                    text = "".join(text_parts)

                    # Get URL
                    r_id = elem.get(qn("r:id"))
                    anchor = elem.get(qn("w:anchor"))

                    if anchor:
                        url = f"#{anchor}"
                    elif r_id:
                        try:
                            rel = self._document.part.rels[r_id]
                            url = rel.target_ref
                        except KeyError:
                            url = ""
                    else:
                        url = ""

                    if text:
                        hyperlinks.append(
                            HyperlinkDTO(
                                text=text,
                                url=url,
                                paragraph_index=i,
                            )
                        )

        return hyperlinks

    def add_heading(
        self,
        text: str,
        level: int = 1,
    ) -> int:
        """Add a heading to the document.

        Args:
            text: Heading text.
            level: Heading level (1-9).

        Returns:
            Index of the heading paragraph.

        Raises:
            ValidationError: If the level is out of range.
        """
        if level < 1 or level > 9:
            raise ValidationError("Heading level must be between 1 and 9")

        self._document.add_heading(text, level=level)
        return len(self._document.paragraphs) - 1

    def get_headings(self) -> list[dict[str, Any]]:
        """Get all headings in the document.

        Returns:
            List of heading information.
        """
        headings = []

        for i, para in enumerate(self._document.paragraphs):
            style_name = para.style.name if para.style else ""
            if style_name.startswith("Heading"):
                try:
                    level = int(style_name.replace("Heading ", ""))
                except ValueError:
                    level = 1

                headings.append(
                    {
                        "index": i,
                        "text": para.text,
                        "level": level,
                    }
                )

        return headings
