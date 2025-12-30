"""Document handler for DOCX file operations.

This module provides the core document handling functionality including
reading, writing, validation, and metadata extraction.
"""

import io
from pathlib import Path
from typing import Any, BinaryIO

from docx import Document

from src.core.constants import SUPPORTED_FORMATS
from src.core.exceptions import InvalidDocumentError, UnsupportedFormatError
from src.models.dto import DocumentMetadataDTO


class DocumentHandler:
    """Handler for DOCX document operations.

    This class provides methods for reading, writing, validating,
    and extracting information from DOCX files.
    """

    def __init__(self) -> None:
        """Initialize the document handler."""
        self._document: Document | None = None
        self._file_path: str | None = None

    @property
    def document(self) -> Document:
        """Get the current document.

        Returns:
            The current Document instance.

        Raises:
            InvalidDocumentError: If no document is loaded.
        """
        if self._document is None:
            raise InvalidDocumentError("No document loaded")
        return self._document

    def create_document(self) -> Document:
        """Create a new empty DOCX document.

        Returns:
            A new Document instance.
        """
        self._document = Document()
        self._file_path = None
        return self._document

    def open_document(self, file_path: str | Path) -> Document:
        """Open a DOCX document from a file path.

        Args:
            file_path: Path to the DOCX file.

        Returns:
            The opened Document instance.

        Raises:
            InvalidDocumentError: If the file cannot be opened.
            UnsupportedFormatError: If the file format is not supported.
        """
        path = Path(file_path)
        if not path.exists():
            raise InvalidDocumentError(f"File not found: {file_path}")

        ext = path.suffix.lower()
        if ext not in SUPPORTED_FORMATS:
            raise UnsupportedFormatError(
                f"Unsupported format: {ext}",
                format_=ext,
                supported_formats=list(SUPPORTED_FORMATS.keys()),
            )

        try:
            self._document = Document(str(path))
            self._file_path = str(path)
            return self._document
        except Exception as e:
            raise InvalidDocumentError(f"Failed to open document: {e}")

    def open_from_bytes(self, content: bytes) -> Document:
        """Open a DOCX document from bytes.

        Args:
            content: Document content as bytes.

        Returns:
            The opened Document instance.

        Raises:
            InvalidDocumentError: If the content cannot be parsed.
        """
        try:
            self._document = Document(io.BytesIO(content))
            self._file_path = None
            return self._document
        except Exception as e:
            raise InvalidDocumentError(f"Failed to parse document: {e}")

    def open_from_stream(self, stream: BinaryIO) -> Document:
        """Open a DOCX document from a file-like object.

        Args:
            stream: File-like object containing document data.

        Returns:
            The opened Document instance.

        Raises:
            InvalidDocumentError: If the stream cannot be parsed.
        """
        try:
            self._document = Document(stream)
            self._file_path = None
            return self._document
        except Exception as e:
            raise InvalidDocumentError(f"Failed to parse document stream: {e}")

    def save_document(self, file_path: str | Path | None = None) -> str:
        """Save the document to a file.

        Args:
            file_path: Path to save the document. If None, saves to original path.

        Returns:
            The path where the document was saved.

        Raises:
            InvalidDocumentError: If no document is loaded or no path provided.
        """
        if self._document is None:
            raise InvalidDocumentError("No document to save")

        save_path = file_path or self._file_path
        if save_path is None:
            raise InvalidDocumentError("No file path provided")

        path = Path(save_path)
        path.parent.mkdir(parents=True, exist_ok=True)
        self._document.save(str(path))
        self._file_path = str(path)
        return self._file_path

    def save_to_bytes(self) -> bytes:
        """Save the document to bytes.

        Returns:
            Document content as bytes.

        Raises:
            InvalidDocumentError: If no document is loaded.
        """
        if self._document is None:
            raise InvalidDocumentError("No document to save")

        buffer = io.BytesIO()
        self._document.save(buffer)
        buffer.seek(0)
        return buffer.read()

    def validate_document(self, file_path: str | Path) -> bool:
        """Validate if a file is a valid DOCX document.

        Args:
            file_path: Path to the file to validate.

        Returns:
            True if the file is a valid DOCX document.

        Raises:
            InvalidDocumentError: If validation fails.
        """
        try:
            Document(str(file_path))
            return True
        except Exception as e:
            raise InvalidDocumentError(f"Invalid document: {e}")

    def validate_bytes(self, content: bytes) -> bool:
        """Validate if bytes represent a valid DOCX document.

        Args:
            content: Document content as bytes.

        Returns:
            True if the content is a valid DOCX document.

        Raises:
            InvalidDocumentError: If validation fails.
        """
        try:
            Document(io.BytesIO(content))
            return True
        except Exception as e:
            raise InvalidDocumentError(f"Invalid document content: {e}")

    def get_metadata(self) -> DocumentMetadataDTO:
        """Extract metadata from the current document.

        Returns:
            Document metadata DTO.

        Raises:
            InvalidDocumentError: If no document is loaded.
        """
        doc = self.document
        core_props = doc.core_properties

        return DocumentMetadataDTO(
            author=core_props.author,
            title=core_props.title,
            subject=core_props.subject,
            keywords=core_props.keywords,
            comments=core_props.comments,
            category=core_props.category,
            created=core_props.created,
            modified=core_props.modified,
            last_modified_by=core_props.last_modified_by,
            revision=core_props.revision,
        )

    def set_metadata(
        self,
        author: str | None = None,
        title: str | None = None,
        subject: str | None = None,
        keywords: str | None = None,
        comments: str | None = None,
        category: str | None = None,
    ) -> None:
        """Set document metadata.

        Args:
            author: Document author.
            title: Document title.
            subject: Document subject.
            keywords: Document keywords.
            comments: Document comments.
            category: Document category.

        Raises:
            InvalidDocumentError: If no document is loaded.
        """
        doc = self.document
        core_props = doc.core_properties

        if author is not None:
            core_props.author = author
        if title is not None:
            core_props.title = title
        if subject is not None:
            core_props.subject = subject
        if keywords is not None:
            core_props.keywords = keywords
        if comments is not None:
            core_props.comments = comments
        if category is not None:
            core_props.category = category

    def get_paragraph_count(self) -> int:
        """Get the number of paragraphs in the document.

        Returns:
            Number of paragraphs.

        Raises:
            InvalidDocumentError: If no document is loaded.
        """
        return len(self.document.paragraphs)

    def get_table_count(self) -> int:
        """Get the number of tables in the document.

        Returns:
            Number of tables.

        Raises:
            InvalidDocumentError: If no document is loaded.
        """
        return len(self.document.tables)

    def get_section_count(self) -> int:
        """Get the number of sections in the document.

        Returns:
            Number of sections.

        Raises:
            InvalidDocumentError: If no document is loaded.
        """
        return len(self.document.sections)

    def get_document_structure(self) -> dict[str, Any]:
        """Get a summary of the document structure.

        Returns:
            Dictionary with document structure information.

        Raises:
            InvalidDocumentError: If no document is loaded.
        """
        doc = self.document
        return {
            "paragraphs": len(doc.paragraphs),
            "tables": len(doc.tables),
            "sections": len(doc.sections),
            "styles": len(doc.styles),
            "inline_shapes": len(doc.inline_shapes),
        }

    def get_all_text(self) -> str:
        """Get all text content from the document.

        Returns:
            Concatenated text from all paragraphs.

        Raises:
            InvalidDocumentError: If no document is loaded.
        """
        return "\n".join(para.text for para in self.document.paragraphs)

    def get_word_count(self) -> int:
        """Get the word count of the document.

        Returns:
            Number of words in the document.

        Raises:
            InvalidDocumentError: If no document is loaded.
        """
        text = self.get_all_text()
        return len(text.split())

    def get_character_count(self, include_spaces: bool = True) -> int:
        """Get the character count of the document.

        Args:
            include_spaces: Whether to include spaces in the count.

        Returns:
            Number of characters in the document.

        Raises:
            InvalidDocumentError: If no document is loaded.
        """
        text = self.get_all_text()
        if include_spaces:
            return len(text)
        return len(text.replace(" ", "").replace("\n", ""))

    def close(self) -> None:
        """Close the current document and release resources."""
        self._document = None
        self._file_path = None
