"""Table handler for table operations.

This module provides functionality for creating, reading, and modifying
tables in DOCX documents.
"""

from typing import Any

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Inches, Pt
from docx.table import Table, _Cell, _Row

from src.core.constants import MAX_TABLE_COLUMNS, MAX_TABLE_ROWS
from src.core.enums import TableBorderStyle
from src.core.exceptions import ValidationError
from src.models.dto import CellDTO, TableDTO


class TableHandler:
    """Handler for table operations.

    This class provides methods for creating, reading, and modifying
    tables in DOCX documents.
    """

    def __init__(self, document: Document) -> None:
        """Initialize the table handler.

        Args:
            document: The Document instance to work with.
        """
        self._document = document

    @property
    def document(self) -> Document:
        """Get the document instance."""
        return self._document

    def get_table(self, index: int) -> TableDTO:
        """Get a table by index.

        Args:
            index: Table index (0-based).

        Returns:
            Table DTO with structure and content.

        Raises:
            ValidationError: If the index is out of range.
        """
        self._validate_table_index(index)
        table = self._document.tables[index]

        rows = len(table.rows)
        cols = len(table.columns)

        cells = []
        for row_idx, row in enumerate(table.rows):
            row_cells = []
            for col_idx, cell in enumerate(row.cells):
                row_cells.append(
                    CellDTO(
                        row=row_idx,
                        col=col_idx,
                        text=cell.text,
                    )
                )
            cells.append(row_cells)

        return TableDTO(
            index=index,
            rows=rows,
            cols=cols,
            cells=cells,
            style=table.style.name if table.style else None,
        )

    def get_all_tables(self) -> list[TableDTO]:
        """Get all tables in the document.

        Returns:
            List of table DTOs.
        """
        return [
            self.get_table(i)
            for i in range(len(self._document.tables))
        ]

    def add_table(
        self,
        rows: int,
        cols: int,
        style: str | None = None,
        data: list[list[str]] | None = None,
    ) -> int:
        """Add a new table to the document.

        Args:
            rows: Number of rows.
            cols: Number of columns.
            style: Optional table style name.
            data: Optional 2D list of cell values.

        Returns:
            Index of the new table.

        Raises:
            ValidationError: If rows or cols exceed limits.
        """
        self._validate_table_dimensions(rows, cols)

        table = self._document.add_table(rows=rows, cols=cols)

        if style:
            table.style = style

        if data:
            self._fill_table_data(table, data)

        return len(self._document.tables) - 1

    def insert_table_after_paragraph(
        self,
        paragraph_index: int,
        rows: int,
        cols: int,
        style: str | None = None,
        data: list[list[str]] | None = None,
    ) -> int:
        """Insert a table after a specific paragraph.

        Args:
            paragraph_index: Index of the paragraph after which to insert.
            rows: Number of rows.
            cols: Number of columns.
            style: Optional table style name.
            data: Optional 2D list of cell values.

        Returns:
            Index of the new table.

        Raises:
            ValidationError: If indices are out of range.
        """
        self._validate_table_dimensions(rows, cols)

        if paragraph_index < 0 or paragraph_index >= len(self._document.paragraphs):
            raise ValidationError(
                f"Paragraph index {paragraph_index} out of range"
            )

        # Create table and insert after the paragraph
        table = self._document.add_table(rows=rows, cols=cols)

        if style:
            table.style = style

        if data:
            self._fill_table_data(table, data)

        # Move table after the specified paragraph
        para = self._document.paragraphs[paragraph_index]
        para._element.addnext(table._tbl)

        return len(self._document.tables) - 1

    def delete_table(self, index: int) -> None:
        """Delete a table by index.

        Args:
            index: Table index to delete.

        Raises:
            ValidationError: If the index is out of range.
        """
        self._validate_table_index(index)
        table = self._document.tables[index]
        tbl = table._tbl
        tbl.getparent().remove(tbl)

    def get_cell(self, table_index: int, row: int, col: int) -> CellDTO:
        """Get a specific cell from a table.

        Args:
            table_index: Index of the table.
            row: Row index (0-based).
            col: Column index (0-based).

        Returns:
            Cell DTO.

        Raises:
            ValidationError: If indices are out of range.
        """
        self._validate_table_index(table_index)
        table = self._document.tables[table_index]
        self._validate_cell_indices(table, row, col)

        cell = table.cell(row, col)
        return CellDTO(
            row=row,
            col=col,
            text=cell.text,
        )

    def set_cell(
        self,
        table_index: int,
        row: int,
        col: int,
        text: str,
    ) -> CellDTO:
        """Set the content of a specific cell.

        Args:
            table_index: Index of the table.
            row: Row index (0-based).
            col: Column index (0-based).
            text: Text content for the cell.

        Returns:
            Updated cell DTO.

        Raises:
            ValidationError: If indices are out of range.
        """
        self._validate_table_index(table_index)
        table = self._document.tables[table_index]
        self._validate_cell_indices(table, row, col)

        cell = table.cell(row, col)
        cell.text = text

        return CellDTO(row=row, col=col, text=text)

    def add_row(self, table_index: int) -> int:
        """Add a new row to a table.

        Args:
            table_index: Index of the table.

        Returns:
            Index of the new row.

        Raises:
            ValidationError: If the index is out of range.
        """
        self._validate_table_index(table_index)
        table = self._document.tables[table_index]

        if len(table.rows) >= MAX_TABLE_ROWS:
            raise ValidationError(f"Maximum rows ({MAX_TABLE_ROWS}) exceeded")

        table.add_row()
        return len(table.rows) - 1

    def add_column(self, table_index: int) -> int:
        """Add a new column to a table.

        Args:
            table_index: Index of the table.

        Returns:
            Index of the new column.

        Raises:
            ValidationError: If the index is out of range.
        """
        self._validate_table_index(table_index)
        table = self._document.tables[table_index]

        if len(table.columns) >= MAX_TABLE_COLUMNS:
            raise ValidationError(f"Maximum columns ({MAX_TABLE_COLUMNS}) exceeded")

        table.add_column(Inches(1.5))
        return len(table.columns) - 1

    def delete_row(self, table_index: int, row_index: int) -> None:
        """Delete a row from a table.

        Args:
            table_index: Index of the table.
            row_index: Index of the row to delete.

        Raises:
            ValidationError: If indices are out of range.
        """
        self._validate_table_index(table_index)
        table = self._document.tables[table_index]

        if row_index < 0 or row_index >= len(table.rows):
            raise ValidationError(f"Row index {row_index} out of range")

        row = table.rows[row_index]
        tr = row._tr
        tr.getparent().remove(tr)

    def delete_column(self, table_index: int, col_index: int) -> None:
        """Delete a column from a table.

        Args:
            table_index: Index of the table.
            col_index: Index of the column to delete.

        Raises:
            ValidationError: If indices are out of range.
        """
        self._validate_table_index(table_index)
        table = self._document.tables[table_index]

        if col_index < 0 or col_index >= len(table.columns):
            raise ValidationError(f"Column index {col_index} out of range")

        for row in table.rows:
            cell = row.cells[col_index]
            tc = cell._tc
            tc.getparent().remove(tc)

    def merge_cells(
        self,
        table_index: int,
        start_row: int,
        start_col: int,
        end_row: int,
        end_col: int,
    ) -> None:
        """Merge a range of cells.

        Args:
            table_index: Index of the table.
            start_row: Starting row index.
            start_col: Starting column index.
            end_row: Ending row index.
            end_col: Ending column index.

        Raises:
            ValidationError: If indices are out of range.
        """
        self._validate_table_index(table_index)
        table = self._document.tables[table_index]
        self._validate_cell_indices(table, start_row, start_col)
        self._validate_cell_indices(table, end_row, end_col)

        if start_row > end_row or start_col > end_col:
            raise ValidationError("Invalid cell range for merge")

        start_cell = table.cell(start_row, start_col)
        end_cell = table.cell(end_row, end_col)
        start_cell.merge(end_cell)

    def set_table_style(self, table_index: int, style: str) -> None:
        """Set the style of a table.

        Args:
            table_index: Index of the table.
            style: Style name to apply.

        Raises:
            ValidationError: If the index is out of range.
        """
        self._validate_table_index(table_index)
        table = self._document.tables[table_index]
        table.style = style

    def set_column_width(
        self,
        table_index: int,
        col_index: int,
        width_inches: float,
    ) -> None:
        """Set the width of a column.

        Args:
            table_index: Index of the table.
            col_index: Index of the column.
            width_inches: Width in inches.

        Raises:
            ValidationError: If indices are out of range.
        """
        self._validate_table_index(table_index)
        table = self._document.tables[table_index]

        if col_index < 0 or col_index >= len(table.columns):
            raise ValidationError(f"Column index {col_index} out of range")

        for row in table.rows:
            row.cells[col_index].width = Inches(width_inches)

    def get_table_as_list(self, table_index: int) -> list[list[str]]:
        """Get table content as a 2D list.

        Args:
            table_index: Index of the table.

        Returns:
            2D list of cell text values.

        Raises:
            ValidationError: If the index is out of range.
        """
        self._validate_table_index(table_index)
        table = self._document.tables[table_index]

        return [
            [cell.text for cell in row.cells]
            for row in table.rows
        ]

    def _validate_table_index(self, index: int) -> None:
        """Validate that a table index is in range.

        Args:
            index: Table index to validate.

        Raises:
            ValidationError: If the index is out of range.
        """
        if index < 0 or index >= len(self._document.tables):
            raise ValidationError(
                f"Table index {index} out of range (0-{len(self._document.tables) - 1})"
            )

    def _validate_table_dimensions(self, rows: int, cols: int) -> None:
        """Validate table dimensions.

        Args:
            rows: Number of rows.
            cols: Number of columns.

        Raises:
            ValidationError: If dimensions exceed limits.
        """
        if rows < 1 or rows > MAX_TABLE_ROWS:
            raise ValidationError(
                f"Rows must be between 1 and {MAX_TABLE_ROWS}"
            )
        if cols < 1 or cols > MAX_TABLE_COLUMNS:
            raise ValidationError(
                f"Columns must be between 1 and {MAX_TABLE_COLUMNS}"
            )

    def _validate_cell_indices(self, table: Table, row: int, col: int) -> None:
        """Validate cell indices for a table.

        Args:
            table: The table to validate against.
            row: Row index.
            col: Column index.

        Raises:
            ValidationError: If indices are out of range.
        """
        if row < 0 or row >= len(table.rows):
            raise ValidationError(
                f"Row index {row} out of range (0-{len(table.rows) - 1})"
            )
        if col < 0 or col >= len(table.columns):
            raise ValidationError(
                f"Column index {col} out of range (0-{len(table.columns) - 1})"
            )

    def _fill_table_data(self, table: Table, data: list[list[str]]) -> None:
        """Fill table cells with data.

        Args:
            table: The table to fill.
            data: 2D list of cell values.
        """
        for row_idx, row_data in enumerate(data):
            if row_idx >= len(table.rows):
                break
            for col_idx, cell_value in enumerate(row_data):
                if col_idx >= len(table.columns):
                    break
                table.cell(row_idx, col_idx).text = str(cell_value)
