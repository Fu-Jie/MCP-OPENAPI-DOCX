"""Revision handler for track changes operations.

This module provides functionality for managing revisions
and track changes in DOCX documents.
"""

from datetime import datetime
from typing import Any

from docx import Document

from src.core.enums import RevisionAction
from src.core.exceptions import ValidationError


class RevisionHandler:
    """Handler for revision tracking operations.

    This class provides methods for managing revisions
    and track changes in DOCX documents.
    """

    def __init__(self, document: Document) -> None:
        """Initialize the revision handler.

        Args:
            document: The Document instance to work with.
        """
        self._document = document
        self._revisions: list[dict[str, Any]] = []
        self._next_id = 0
        self._tracking_enabled = False

    @property
    def document(self) -> Document:
        """Get the document instance."""
        return self._document

    @property
    def tracking_enabled(self) -> bool:
        """Check if revision tracking is enabled."""
        return self._tracking_enabled

    def enable_tracking(self) -> None:
        """Enable revision tracking."""
        self._tracking_enabled = True

    def disable_tracking(self) -> None:
        """Disable revision tracking."""
        self._tracking_enabled = False

    def add_revision(
        self,
        action: RevisionAction,
        author: str,
        paragraph_index: int,
        original_content: str | None = None,
        new_content: str | None = None,
    ) -> dict[str, Any]:
        """Record a revision.

        Args:
            action: Type of revision action.
            author: Author of the revision.
            paragraph_index: Index of the affected paragraph.
            original_content: Original content before change.
            new_content: New content after change.

        Returns:
            Revision information dictionary.

        Raises:
            ValidationError: If the paragraph index is out of range.
        """
        if paragraph_index < 0 or paragraph_index >= len(self._document.paragraphs):
            raise ValidationError(f"Paragraph index {paragraph_index} out of range")

        revision_id = self._next_id
        self._next_id += 1

        revision = {
            "id": revision_id,
            "action": action,
            "author": author,
            "paragraph_index": paragraph_index,
            "original_content": original_content,
            "new_content": new_content,
            "is_accepted": False,
            "is_rejected": False,
            "created_at": datetime.now(),
            "accepted_at": None,
            "accepted_by": None,
        }

        self._revisions.append(revision)
        return revision

    def get_revision(self, revision_id: int) -> dict[str, Any]:
        """Get a revision by ID.

        Args:
            revision_id: Revision ID.

        Returns:
            Revision information dictionary.

        Raises:
            ValidationError: If the revision is not found.
        """
        for revision in self._revisions:
            if revision["id"] == revision_id:
                return revision

        raise ValidationError(f"Revision not found: {revision_id}")

    def get_all_revisions(self) -> list[dict[str, Any]]:
        """Get all revisions in the document.

        Returns:
            List of revision dictionaries.
        """
        return self._revisions.copy()

    def get_pending_revisions(self) -> list[dict[str, Any]]:
        """Get all pending (not accepted/rejected) revisions.

        Returns:
            List of pending revisions.
        """
        return [
            r for r in self._revisions if not r["is_accepted"] and not r["is_rejected"]
        ]

    def accept_revision(
        self,
        revision_id: int,
        accepted_by: str | None = None,
    ) -> dict[str, Any]:
        """Accept a revision.

        Args:
            revision_id: Revision ID.
            accepted_by: User who accepted the revision.

        Returns:
            Updated revision dictionary.

        Raises:
            ValidationError: If the revision is not found.
        """
        revision = self.get_revision(revision_id)

        if revision["is_accepted"] or revision["is_rejected"]:
            raise ValidationError("Revision has already been processed")

        revision["is_accepted"] = True
        revision["accepted_at"] = datetime.now()
        revision["accepted_by"] = accepted_by

        # Apply the revision to the document
        self._apply_revision(revision)

        return revision

    def reject_revision(
        self,
        revision_id: int,
        rejected_by: str | None = None,
    ) -> dict[str, Any]:
        """Reject a revision.

        Args:
            revision_id: Revision ID.
            rejected_by: User who rejected the revision.

        Returns:
            Updated revision dictionary.

        Raises:
            ValidationError: If the revision is not found.
        """
        revision = self.get_revision(revision_id)

        if revision["is_accepted"] or revision["is_rejected"]:
            raise ValidationError("Revision has already been processed")

        revision["is_rejected"] = True
        revision["accepted_at"] = datetime.now()
        revision["accepted_by"] = rejected_by

        return revision

    def accept_all_revisions(self, accepted_by: str | None = None) -> int:
        """Accept all pending revisions.

        Args:
            accepted_by: User who accepted the revisions.

        Returns:
            Number of revisions accepted.
        """
        count = 0
        for revision in self.get_pending_revisions():
            try:
                self.accept_revision(revision["id"], accepted_by)
                count += 1
            except ValidationError:
                continue
        return count

    def reject_all_revisions(self, rejected_by: str | None = None) -> int:
        """Reject all pending revisions.

        Args:
            rejected_by: User who rejected the revisions.

        Returns:
            Number of revisions rejected.
        """
        count = 0
        for revision in self.get_pending_revisions():
            try:
                self.reject_revision(revision["id"], rejected_by)
                count += 1
            except ValidationError:
                continue
        return count

    def get_revisions_by_author(self, author: str) -> list[dict[str, Any]]:
        """Get all revisions by a specific author.

        Args:
            author: Author name.

        Returns:
            List of revisions by the author.
        """
        return [r for r in self._revisions if r["author"] == author]

    def get_revisions_by_action(
        self,
        action: RevisionAction,
    ) -> list[dict[str, Any]]:
        """Get all revisions of a specific action type.

        Args:
            action: Revision action type.

        Returns:
            List of revisions with the specified action.
        """
        return [r for r in self._revisions if r["action"] == action]

    def get_revision_count(self) -> dict[str, int]:
        """Get revision statistics.

        Returns:
            Dictionary with revision counts.
        """
        pending = len(self.get_pending_revisions())
        accepted = len([r for r in self._revisions if r["is_accepted"]])
        rejected = len([r for r in self._revisions if r["is_rejected"]])

        return {
            "total": len(self._revisions),
            "pending": pending,
            "accepted": accepted,
            "rejected": rejected,
        }

    def compare_paragraphs(
        self,
        index1: int,
        index2: int,
    ) -> dict[str, Any]:
        """Compare two paragraphs.

        Args:
            index1: Index of first paragraph.
            index2: Index of second paragraph.

        Returns:
            Comparison result dictionary.

        Raises:
            ValidationError: If indices are out of range.
        """
        paras = self._document.paragraphs

        if index1 < 0 or index1 >= len(paras):
            raise ValidationError(f"Paragraph index {index1} out of range")
        if index2 < 0 or index2 >= len(paras):
            raise ValidationError(f"Paragraph index {index2} out of range")

        text1 = paras[index1].text
        text2 = paras[index2].text

        return {
            "paragraph1": {"index": index1, "text": text1},
            "paragraph2": {"index": index2, "text": text2},
            "are_identical": text1 == text2,
            "length_diff": len(text2) - len(text1),
        }

    def _apply_revision(self, revision: dict[str, Any]) -> None:
        """Apply an accepted revision to the document.

        Args:
            revision: Revision to apply.
        """
        para = self._document.paragraphs[revision["paragraph_index"]]

        action = revision["action"]

        if action == RevisionAction.INSERT:
            if revision["new_content"]:
                para.add_run(revision["new_content"])

        elif action == RevisionAction.DELETE:
            # Clear the paragraph content
            para.clear()

        elif action == RevisionAction.REPLACE:
            if revision["new_content"]:
                para.clear()
                para.add_run(revision["new_content"])

    def export_revisions(self) -> list[dict[str, Any]]:
        """Export all revisions for external processing.

        Returns:
            List of revision dictionaries with all details.
        """
        exported = []
        for revision in self._revisions:
            exported.append(
                {
                    "id": revision["id"],
                    "action": (
                        revision["action"].value
                        if isinstance(revision["action"], RevisionAction)
                        else revision["action"]
                    ),
                    "author": revision["author"],
                    "paragraph_index": revision["paragraph_index"],
                    "original_content": revision["original_content"],
                    "new_content": revision["new_content"],
                    "is_accepted": revision["is_accepted"],
                    "is_rejected": revision["is_rejected"],
                    "created_at": revision["created_at"].isoformat(),
                    "accepted_at": (
                        revision["accepted_at"].isoformat()
                        if revision["accepted_at"]
                        else None
                    ),
                    "accepted_by": revision["accepted_by"],
                }
            )
        return exported

    def clear_revision_history(self) -> int:
        """Clear all revision history.

        Returns:
            Number of revisions cleared.
        """
        count = len(self._revisions)
        self._revisions.clear()
        self._next_id = 0
        return count
