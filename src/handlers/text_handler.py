"""Text handler for paragraph and text operations.

This module provides functionality for manipulating text content
in DOCX documents including paragraphs, runs, and formatting.
"""

from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from src.core.enums import TextAlignment
from src.core.exceptions import ValidationError
from src.models.dto import ParagraphDTO, RunDTO
from src.models.schemas import TextFormat


class TextHandler:
    """Handler for text and paragraph operations.

    This class provides methods for reading, creating, and modifying
    text content in DOCX documents.
    """

    ALIGNMENT_MAP = {
        TextAlignment.LEFT: WD_ALIGN_PARAGRAPH.LEFT,
        TextAlignment.CENTER: WD_ALIGN_PARAGRAPH.CENTER,
        TextAlignment.RIGHT: WD_ALIGN_PARAGRAPH.RIGHT,
        TextAlignment.JUSTIFY: WD_ALIGN_PARAGRAPH.JUSTIFY,
        TextAlignment.DISTRIBUTE: WD_ALIGN_PARAGRAPH.DISTRIBUTE,
    }

    def __init__(self, document: Document | None = None) -> None:
        """Initialize the text handler.

        Args:
            document: The Document instance to work with (optional).
        """
        self._document = document

    @property
    def document(self) -> Document:
        """Get the document instance."""
        if self._document is None:
            raise ValueError("No document loaded")
        return self._document

    def set_document(self, document: Document) -> None:
        """Set the document instance.

        Args:
            document: The Document instance to work with.
        """
        self._document = document

    def get_paragraph(self, index: int) -> ParagraphDTO:
        """Get a paragraph by index.

        Args:
            index: Paragraph index (0-based).

        Returns:
            Paragraph DTO with text and formatting information.

        Raises:
            ValidationError: If the index is out of range.
        """
        self._validate_paragraph_index(index)
        para = self._document.paragraphs[index]

        runs = []
        for run in para.runs:
            runs.append(
                RunDTO(
                    text=run.text,
                    bold=run.bold or False,
                    italic=run.italic or False,
                    underline=run.underline or False,
                    font_name=run.font.name,
                    font_size=int(run.font.size.pt) if run.font.size else None,
                    color=(
                        self._get_color_hex(run.font.color.rgb)
                        if run.font.color.rgb
                        else None
                    ),
                )
            )

        alignment = None
        if para.alignment is not None:
            for key, value in self.ALIGNMENT_MAP.items():
                if para.alignment == value:
                    alignment = key
                    break

        return ParagraphDTO(
            index=index,
            text=para.text,
            style=para.style.name if para.style else None,
            alignment=alignment,
            runs=runs,
        )

    def get_all_paragraphs(self) -> list[ParagraphDTO]:
        """Get all paragraphs in the document.

        Returns:
            List of paragraph DTOs.
        """
        return [self.get_paragraph(i) for i in range(len(self._document.paragraphs))]

    def add_paragraph(
        self,
        text: str,
        style: str | None = None,
        alignment: TextAlignment | None = None,
    ) -> int:
        """Add a new paragraph to the document.

        Args:
            text: Paragraph text content.
            style: Optional style name to apply.
            alignment: Optional text alignment.

        Returns:
            Index of the new paragraph.
        """
        para = self._document.add_paragraph(text, style=style)

        if alignment is not None:
            para.alignment = self.ALIGNMENT_MAP.get(alignment)

        return len(self._document.paragraphs) - 1

    def insert_paragraph(
        self,
        index: int,
        text: str,
        style: str | None = None,
        alignment: TextAlignment | None = None,
    ) -> int:
        """Insert a paragraph at a specific index.

        Args:
            index: Index where to insert the paragraph.
            text: Paragraph text content.
            style: Optional style name to apply.
            alignment: Optional text alignment.

        Returns:
            Index of the inserted paragraph.

        Raises:
            ValidationError: If the index is out of range.
        """
        if index < 0 or index > len(self._document.paragraphs):
            raise ValidationError(
                f"Index {index} out of range (0-{len(self._document.paragraphs)})"
            )

        if index == len(self._document.paragraphs):
            return self.add_paragraph(text, style, alignment)

        # Get the paragraph at the target index
        target_para = self._document.paragraphs[index]

        # Create a new paragraph element before the target
        new_para = target_para._element.makeelement(qn("w:p"), {})
        target_para._element.addprevious(new_para)

        # Now get the new paragraph through the document
        para = self._document.paragraphs[index]
        para.add_run(text)

        if style:
            para.style = style
        if alignment:
            para.alignment = self.ALIGNMENT_MAP.get(alignment)

        return index

    def update_paragraph(
        self,
        index: int,
        text: str | None = None,
        style: str | None = None,
        alignment: TextAlignment | None = None,
    ) -> ParagraphDTO:
        """Update an existing paragraph.

        Args:
            index: Paragraph index.
            text: New text content (if provided).
            style: New style name (if provided).
            alignment: New alignment (if provided).

        Returns:
            Updated paragraph DTO.

        Raises:
            ValidationError: If the index is out of range.
        """
        self._validate_paragraph_index(index)
        para = self._document.paragraphs[index]

        if text is not None:
            para.clear()
            para.add_run(text)

        if style is not None:
            para.style = style

        if alignment is not None:
            para.alignment = self.ALIGNMENT_MAP.get(alignment)

        return self.get_paragraph(index)

    def delete_paragraph(self, index: int) -> None:
        """Delete a paragraph by index.

        Args:
            index: Paragraph index to delete.

        Raises:
            ValidationError: If the index is out of range.
        """
        self._validate_paragraph_index(index)
        para = self._document.paragraphs[index]
        p = para._element
        p.getparent().remove(p)

    def add_run(
        self,
        paragraph_index: int,
        text: str,
        format_: TextFormat | None = None,
    ) -> RunDTO:
        """Add a run of text to a paragraph.

        Args:
            paragraph_index: Index of the paragraph.
            text: Text content to add.
            format_: Optional text formatting.

        Returns:
            The created run DTO.

        Raises:
            ValidationError: If the index is out of range.
        """
        self._validate_paragraph_index(paragraph_index)
        para = self._document.paragraphs[paragraph_index]
        run = para.add_run(text)

        if format_:
            self._apply_format(run, format_)

        return RunDTO(
            text=run.text,
            bold=run.bold or False,
            italic=run.italic or False,
            underline=run.underline or False,
            font_name=run.font.name,
            font_size=int(run.font.size.pt) if run.font.size else None,
        )

    def format_run(
        self,
        paragraph_index: int,
        run_index: int,
        format_: TextFormat,
    ) -> RunDTO:
        """Apply formatting to a specific run.

        Args:
            paragraph_index: Index of the paragraph.
            run_index: Index of the run within the paragraph.
            format_: Text formatting to apply.

        Returns:
            Updated run DTO.

        Raises:
            ValidationError: If the indices are out of range.
        """
        self._validate_paragraph_index(paragraph_index)
        para = self._document.paragraphs[paragraph_index]

        if run_index < 0 or run_index >= len(para.runs):
            raise ValidationError(
                f"Run index {run_index} out of range (0-{len(para.runs) - 1})"
            )

        run = para.runs[run_index]
        self._apply_format(run, format_)

        return RunDTO(
            text=run.text,
            bold=run.bold or False,
            italic=run.italic or False,
            underline=run.underline or False,
            font_name=run.font.name,
            font_size=int(run.font.size.pt) if run.font.size else None,
        )

    def insert_text(
        self,
        paragraph_index: int,
        offset: int,
        text: str,
        format_: TextFormat | None = None,
    ) -> None:
        """Insert text at a specific position within a paragraph.

        Args:
            paragraph_index: Index of the paragraph.
            offset: Character offset where to insert.
            text: Text to insert.
            format_: Optional formatting for the inserted text.

        Raises:
            ValidationError: If the index or offset is out of range.
        """
        self._validate_paragraph_index(paragraph_index)
        para = self._document.paragraphs[paragraph_index]
        para_text = para.text

        if offset < 0 or offset > len(para_text):
            raise ValidationError(f"Offset {offset} out of range (0-{len(para_text)})")

        # Clear and rebuild with inserted text
        new_text = para_text[:offset] + text + para_text[offset:]
        para.clear()
        run = para.add_run(new_text)

        if format_:
            self._apply_format(run, format_)

    def find_text(
        self,
        search_text: str,
        case_sensitive: bool = False,
        whole_word: bool = False,
    ) -> list[dict[str, Any]]:
        """Find text occurrences in the document.

        Args:
            search_text: Text to search for.
            case_sensitive: Whether the search is case-sensitive.
            whole_word: Whether to match whole words only.

        Returns:
            List of matches with paragraph index and offset.
        """
        results = []
        search = search_text if case_sensitive else search_text.lower()

        for i, para in enumerate(self._document.paragraphs):
            text = para.text if case_sensitive else para.text.lower()
            start = 0

            while True:
                pos = text.find(search, start)
                if pos == -1:
                    break

                if whole_word:
                    before = pos == 0 or not text[pos - 1].isalnum()
                    after = (
                        pos + len(search) >= len(text)
                        or not text[pos + len(search)].isalnum()
                    )
                    if not (before and after):
                        start = pos + 1
                        continue

                results.append(
                    {
                        "paragraph_index": i,
                        "offset": pos,
                        "length": len(search_text),
                        "text": para.text[pos : pos + len(search_text)],
                    }
                )
                start = pos + 1

        return results

    def replace_text(
        self,
        find: str,
        replace: str,
        case_sensitive: bool = False,
        whole_word: bool = False,
    ) -> int:
        """Replace text throughout the document.

        Args:
            find: Text to find.
            replace: Replacement text.
            case_sensitive: Whether the search is case-sensitive.
            whole_word: Whether to match whole words only.

        Returns:
            Number of replacements made.
        """
        count = 0

        for para in self._document.paragraphs:
            for run in para.runs:
                if case_sensitive:
                    if find in run.text:
                        run.text = run.text.replace(find, replace)
                        count += 1
                else:
                    lower_text = run.text.lower()
                    lower_find = find.lower()
                    if lower_find in lower_text:
                        # Case-insensitive replacement
                        import re

                        pattern = re.escape(find)
                        if whole_word:
                            pattern = r"\b" + pattern + r"\b"
                        flags = 0 if case_sensitive else re.IGNORECASE
                        new_text, n = re.subn(pattern, replace, run.text, flags=flags)
                        run.text = new_text
                        count += n

        return count

    def _validate_paragraph_index(self, index: int) -> None:
        """Validate that a paragraph index is in range.

        Args:
            index: Paragraph index to validate.

        Raises:
            ValidationError: If the index is out of range.
        """
        if index < 0 or index >= len(self._document.paragraphs):
            raise ValidationError(
                f"Paragraph index {index} out of range (0-{len(self._document.paragraphs) - 1})"
            )

    def _apply_format(self, run: Any, format_: TextFormat) -> None:
        """Apply formatting to a run.

        Args:
            run: The run to format.
            format_: Formatting options to apply.
        """
        if format_.bold is not None:
            run.bold = format_.bold
        if format_.italic is not None:
            run.italic = format_.italic
        if format_.underline is not None:
            run.underline = format_.underline
        if format_.strike is not None:
            run.font.strike = format_.strike
        if format_.font_name is not None:
            run.font.name = format_.font_name
        if format_.font_size is not None:
            run.font.size = Pt(format_.font_size)
        if format_.color is not None:
            run.font.color.rgb = RGBColor.from_string(format_.color)
        if format_.superscript is not None:
            run.font.superscript = format_.superscript
        if format_.subscript is not None:
            run.font.subscript = format_.subscript

    def _get_color_hex(self, rgb: RGBColor | None) -> str | None:
        """Convert RGBColor to hex string.

        Args:
            rgb: RGBColor instance.

        Returns:
            Hex color string or None.
        """
        if rgb is None:
            return None
        return f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
