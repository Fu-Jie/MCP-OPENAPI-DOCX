"""Style handler for document style operations.

This module provides functionality for managing and applying
styles in DOCX documents.
"""

import contextlib
from typing import Any

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from src.core.enums import StyleType, TextAlignment
from src.core.exceptions import ValidationError
from src.models.dto import StyleDTO
from src.models.schemas import StyleCreate


class StyleHandler:
    """Handler for style operations.

    This class provides methods for managing and applying
    document styles.
    """

    STYLE_TYPE_MAP = {
        StyleType.PARAGRAPH: WD_STYLE_TYPE.PARAGRAPH,
        StyleType.CHARACTER: WD_STYLE_TYPE.CHARACTER,
        StyleType.TABLE: WD_STYLE_TYPE.TABLE,
        StyleType.NUMBERING: WD_STYLE_TYPE.LIST,
    }

    ALIGNMENT_MAP = {
        TextAlignment.LEFT: WD_ALIGN_PARAGRAPH.LEFT,
        TextAlignment.CENTER: WD_ALIGN_PARAGRAPH.CENTER,
        TextAlignment.RIGHT: WD_ALIGN_PARAGRAPH.RIGHT,
        TextAlignment.JUSTIFY: WD_ALIGN_PARAGRAPH.JUSTIFY,
        TextAlignment.DISTRIBUTE: WD_ALIGN_PARAGRAPH.DISTRIBUTE,
    }

    def __init__(self, document: Document) -> None:
        """Initialize the style handler.

        Args:
            document: The Document instance to work with.
        """
        self._document = document

    @property
    def document(self) -> Document:
        """Get the document instance."""
        return self._document

    def get_style(self, name: str) -> StyleDTO:
        """Get a style by name.

        Args:
            name: Style name.

        Returns:
            Style DTO with style information.

        Raises:
            ValidationError: If the style is not found.
        """
        try:
            style = self._document.styles[name]
        except KeyError:
            raise ValidationError(f"Style not found: {name}")

        # Determine style type
        style_type = "paragraph"
        if style.type == WD_STYLE_TYPE.CHARACTER:
            style_type = "character"
        elif style.type == WD_STYLE_TYPE.TABLE:
            style_type = "table"
        elif style.type == WD_STYLE_TYPE.LIST:
            style_type = "numbering"

        # Get font properties
        font_name = None
        font_size = None
        bold = None
        italic = None
        color = None

        if hasattr(style, "font") and style.font:
            font_name = style.font.name
            font_size = int(style.font.size.pt) if style.font.size else None
            bold = style.font.bold
            italic = style.font.italic
            if style.font.color and style.font.color.rgb:
                rgb = style.font.color.rgb
                color = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"

        return StyleDTO(
            name=style.name,
            style_type=style_type,
            base_style=style.base_style.name if style.base_style else None,
            font_name=font_name,
            font_size=font_size,
            bold=bold,
            italic=italic,
            color=color,
        )

    def get_all_styles(
        self,
        style_type: StyleType | None = None,
    ) -> list[StyleDTO]:
        """Get all styles in the document.

        Args:
            style_type: Optional filter by style type.

        Returns:
            List of style DTOs.
        """
        styles = []

        for style in self._document.styles:
            # Skip hidden styles
            if style.hidden:
                continue

            # Filter by type if specified
            if style_type:
                target_type = self.STYLE_TYPE_MAP.get(style_type)
                if style.type != target_type:
                    continue

            try:
                styles.append(self.get_style(style.name))
            except Exception:
                continue

        return styles

    def create_style(self, style_data: StyleCreate) -> StyleDTO:
        """Create a new style.

        Args:
            style_data: Style creation data.

        Returns:
            Created style DTO.

        Raises:
            ValidationError: If a style with the same name exists.
        """
        # Check if style already exists
        try:
            self._document.styles[style_data.name]
            raise ValidationError(f"Style already exists: {style_data.name}")
        except KeyError:
            pass

        # Create the style
        style = self._document.styles.add_style(
            style_data.name,
            WD_STYLE_TYPE.PARAGRAPH,
        )

        # Set base style
        if style_data.base_style:
            with contextlib.suppress(KeyError):
                style.base_style = self._document.styles[style_data.base_style]

        # Apply font properties
        if style_data.font_name:
            style.font.name = style_data.font_name
        if style_data.font_size:
            style.font.size = Pt(style_data.font_size)
        if style_data.bold is not None:
            style.font.bold = style_data.bold
        if style_data.italic is not None:
            style.font.italic = style_data.italic
        if style_data.color:
            style.font.color.rgb = RGBColor.from_string(style_data.color)

        # Apply paragraph properties
        if style_data.alignment:
            style.paragraph_format.alignment = self.ALIGNMENT_MAP.get(
                style_data.alignment
            )
        if style_data.line_spacing:
            style.paragraph_format.line_spacing = style_data.line_spacing
        if style_data.space_before is not None:
            style.paragraph_format.space_before = Pt(style_data.space_before * 12)
        if style_data.space_after is not None:
            style.paragraph_format.space_after = Pt(style_data.space_after * 12)

        return self.get_style(style_data.name)

    def update_style(
        self,
        name: str,
        updates: dict[str, Any],
    ) -> StyleDTO:
        """Update an existing style.

        Args:
            name: Style name.
            updates: Dictionary of updates to apply.

        Returns:
            Updated style DTO.

        Raises:
            ValidationError: If the style is not found.
        """
        try:
            style = self._document.styles[name]
        except KeyError:
            raise ValidationError(f"Style not found: {name}")

        # Apply font updates
        if "font_name" in updates and updates["font_name"]:
            style.font.name = updates["font_name"]
        if "font_size" in updates and updates["font_size"]:
            style.font.size = Pt(updates["font_size"])
        if "bold" in updates:
            style.font.bold = updates["bold"]
        if "italic" in updates:
            style.font.italic = updates["italic"]
        if "color" in updates and updates["color"]:
            style.font.color.rgb = RGBColor.from_string(updates["color"])

        # Apply paragraph updates
        if "alignment" in updates and updates["alignment"]:
            alignment = TextAlignment(updates["alignment"])
            style.paragraph_format.alignment = self.ALIGNMENT_MAP.get(alignment)
        if "line_spacing" in updates and updates["line_spacing"]:
            style.paragraph_format.line_spacing = updates["line_spacing"]

        return self.get_style(name)

    def delete_style(self, name: str) -> None:
        """Delete a custom style.

        Args:
            name: Style name to delete.

        Raises:
            ValidationError: If the style is not found or is built-in.
        """
        try:
            style = self._document.styles[name]
        except KeyError:
            raise ValidationError(f"Style not found: {name}")

        if style.builtin:
            raise ValidationError(f"Cannot delete built-in style: {name}")

        # Remove the style element
        style._element.getparent().remove(style._element)

    def apply_style_to_paragraph(
        self,
        paragraph_index: int,
        style_name: str,
    ) -> None:
        """Apply a style to a paragraph.

        Args:
            paragraph_index: Index of the paragraph.
            style_name: Name of the style to apply.

        Raises:
            ValidationError: If the index or style is invalid.
        """
        if paragraph_index < 0 or paragraph_index >= len(self._document.paragraphs):
            raise ValidationError(
                f"Paragraph index {paragraph_index} out of range"
            )

        try:
            self._document.styles[style_name]
        except KeyError:
            raise ValidationError(f"Style not found: {style_name}")

        para = self._document.paragraphs[paragraph_index]
        para.style = style_name

    def get_paragraph_style(self, paragraph_index: int) -> str | None:
        """Get the style name of a paragraph.

        Args:
            paragraph_index: Index of the paragraph.

        Returns:
            Style name or None.

        Raises:
            ValidationError: If the index is out of range.
        """
        if paragraph_index < 0 or paragraph_index >= len(self._document.paragraphs):
            raise ValidationError(
                f"Paragraph index {paragraph_index} out of range"
            )

        para = self._document.paragraphs[paragraph_index]
        return para.style.name if para.style else None

    def get_built_in_styles(self) -> list[str]:
        """Get a list of built-in style names.

        Returns:
            List of built-in style names.
        """
        return [
            style.name
            for style in self._document.styles
            if style.builtin and not style.hidden
        ]

    def get_custom_styles(self) -> list[str]:
        """Get a list of custom style names.

        Returns:
            List of custom style names.
        """
        return [
            style.name
            for style in self._document.styles
            if not style.builtin and not style.hidden
        ]

    def copy_style(
        self,
        source_name: str,
        new_name: str,
    ) -> StyleDTO:
        """Copy a style with a new name.

        Args:
            source_name: Source style name.
            new_name: New style name.

        Returns:
            Created style DTO.

        Raises:
            ValidationError: If source not found or new name exists.
        """
        source = self.get_style(source_name)

        style_data = StyleCreate(
            name=new_name,
            base_style=source_name,
            font_name=source.font_name,
            font_size=source.font_size,
            bold=source.bold,
            italic=source.italic,
            color=source.color,
        )

        return self.create_style(style_data)
