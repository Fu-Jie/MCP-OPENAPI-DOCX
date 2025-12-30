"""List handler for list and numbering operations.

This module provides functionality for creating and managing
bullet lists, numbered lists, and multi-level lists.
"""

from docx import Document

from src.core.enums import ListType, NumberingFormat
from src.core.exceptions import ValidationError
from src.models.dto import ListDTO, ListItemDTO


class ListHandler:
    """Handler for list operations.

    This class provides methods for creating and managing
    various types of lists in DOCX documents.
    """

    def __init__(self, document: Document | None = None) -> None:
        """Initialize the list handler.

        Args:
            document: The Document instance to work with (optional).
        """
        self._document = document

    @property
    def document(self) -> Document:
        """Get the document instance."""
        if self._document is None:
            raise ValueError("No document loaded")
        return self._document

    def set_document(self, document: Document) -> None:
        """Set the document instance.

        Args:
            document: The Document instance to work with.
        """
        self._document = document

    def create_bullet_list(
        self,
        items: list[str],
        paragraph_index: int | None = None,
    ) -> int:
        """Create a bullet list.

        Args:
            items: List of text items.
            paragraph_index: Optional starting paragraph index.

        Returns:
            Index of the first list paragraph.
        """
        start_index = (
            paragraph_index
            if paragraph_index is not None
            else len(self._document.paragraphs)
        )

        for _i, item in enumerate(items):
            self._document.add_paragraph(item, style="List Bullet")

        return start_index

    def create_numbered_list(
        self,
        items: list[str],
        numbering_format: NumberingFormat = NumberingFormat.DECIMAL,
        paragraph_index: int | None = None,
    ) -> int:
        """Create a numbered list.

        Args:
            items: List of text items.
            numbering_format: Numbering format to use.
            paragraph_index: Optional starting paragraph index.

        Returns:
            Index of the first list paragraph.
        """
        start_index = (
            paragraph_index
            if paragraph_index is not None
            else len(self._document.paragraphs)
        )

        for item in items:
            self._document.add_paragraph(item, style="List Number")

        return start_index

    def create_multilevel_list(
        self,
        items: list[tuple[str, int]],
        paragraph_index: int | None = None,
    ) -> int:
        """Create a multi-level list.

        Args:
            items: List of tuples (text, level).
            paragraph_index: Optional starting paragraph index.

        Returns:
            Index of the first list paragraph.
        """
        start_index = (
            paragraph_index
            if paragraph_index is not None
            else len(self._document.paragraphs)
        )

        for text, level in items:
            if level == 0:
                style = "List Bullet"
            elif level == 1:
                style = "List Bullet 2"
            else:
                style = "List Bullet 3"

            self._document.add_paragraph(text, style=style)

        return start_index

    def add_list_item(
        self,
        text: str,
        list_type: ListType = ListType.BULLET,
        level: int = 0,
    ) -> int:
        """Add a single list item.

        Args:
            text: Item text.
            list_type: Type of list.
            level: Indentation level (0-8).

        Returns:
            Index of the new paragraph.
        """
        if level < 0 or level > 8:
            raise ValidationError("List level must be between 0 and 8")

        if list_type == ListType.BULLET:
            style = f"List Bullet{' ' + str(level + 1) if level > 0 else ''}"
        else:
            style = f"List Number{' ' + str(level + 1) if level > 0 else ''}"

        try:
            self._document.add_paragraph(text, style=style)
        except KeyError:
            # Fall back to basic list style if level-specific style doesn't exist
            if list_type == ListType.BULLET:
                self._document.add_paragraph(text, style="List Bullet")
            else:
                self._document.add_paragraph(text, style="List Number")

        return len(self._document.paragraphs) - 1

    def convert_to_list(
        self,
        paragraph_index: int,
        list_type: ListType = ListType.BULLET,
    ) -> None:
        """Convert a paragraph to a list item.

        Args:
            paragraph_index: Index of the paragraph to convert.
            list_type: Type of list.

        Raises:
            ValidationError: If the index is out of range.
        """
        if paragraph_index < 0 or paragraph_index >= len(self._document.paragraphs):
            raise ValidationError(f"Paragraph index {paragraph_index} out of range")

        para = self._document.paragraphs[paragraph_index]
        if list_type == ListType.BULLET:
            para.style = "List Bullet"
        else:
            para.style = "List Number"

    def remove_list_formatting(self, paragraph_index: int) -> None:
        """Remove list formatting from a paragraph.

        Args:
            paragraph_index: Index of the paragraph.

        Raises:
            ValidationError: If the index is out of range.
        """
        if paragraph_index < 0 or paragraph_index >= len(self._document.paragraphs):
            raise ValidationError(f"Paragraph index {paragraph_index} out of range")

        para = self._document.paragraphs[paragraph_index]
        para.style = "Normal"

    def get_list_items(
        self,
        start_index: int,
        end_index: int | None = None,
    ) -> ListDTO:
        """Get list items from a range of paragraphs.

        Args:
            start_index: Starting paragraph index.
            end_index: Ending paragraph index (exclusive).

        Returns:
            List DTO with items.

        Raises:
            ValidationError: If indices are out of range.
        """
        if start_index < 0 or start_index >= len(self._document.paragraphs):
            raise ValidationError(f"Start index {start_index} out of range")

        if end_index is None:
            end_index = len(self._document.paragraphs)

        items = []
        list_type = ListType.BULLET

        for i in range(start_index, min(end_index, len(self._document.paragraphs))):
            para = self._document.paragraphs[i]
            style_name = para.style.name if para.style else ""

            if "List" not in style_name:
                break

            if "Number" in style_name:
                list_type = ListType.NUMBERED

            # Determine level from style name
            level = 0
            if style_name.endswith("2"):
                level = 1
            elif style_name.endswith("3"):
                level = 2

            items.append(
                ListItemDTO(
                    index=i - start_index,
                    text=para.text,
                    level=level,
                )
            )

        return ListDTO(
            paragraph_index=start_index,
            list_type=list_type,
            items=items,
        )

    def change_list_type(
        self,
        paragraph_index: int,
        list_type: ListType,
    ) -> None:
        """Change the type of a list item.

        Args:
            paragraph_index: Index of the list paragraph.
            list_type: New list type.

        Raises:
            ValidationError: If the index is out of range.
        """
        if paragraph_index < 0 or paragraph_index >= len(self._document.paragraphs):
            raise ValidationError(f"Paragraph index {paragraph_index} out of range")

        para = self._document.paragraphs[paragraph_index]
        current_style = para.style.name if para.style else ""

        # Determine current level
        level = ""
        if current_style.endswith("2"):
            level = " 2"
        elif current_style.endswith("3"):
            level = " 3"

        if list_type == ListType.BULLET:
            para.style = f"List Bullet{level}"
        else:
            para.style = f"List Number{level}"

    def indent_list_item(self, paragraph_index: int) -> None:
        """Increase the indentation level of a list item.

        Args:
            paragraph_index: Index of the list paragraph.

        Raises:
            ValidationError: If the index is out of range.
        """
        if paragraph_index < 0 or paragraph_index >= len(self._document.paragraphs):
            raise ValidationError(f"Paragraph index {paragraph_index} out of range")

        para = self._document.paragraphs[paragraph_index]
        style_name = para.style.name if para.style else ""

        if "List Bullet" in style_name:
            if style_name == "List Bullet":
                para.style = "List Bullet 2"
            elif style_name == "List Bullet 2":
                para.style = "List Bullet 3"
        elif "List Number" in style_name:
            if style_name == "List Number":
                para.style = "List Number 2"
            elif style_name == "List Number 2":
                para.style = "List Number 3"

    def outdent_list_item(self, paragraph_index: int) -> None:
        """Decrease the indentation level of a list item.

        Args:
            paragraph_index: Index of the list paragraph.

        Raises:
            ValidationError: If the index is out of range.
        """
        if paragraph_index < 0 or paragraph_index >= len(self._document.paragraphs):
            raise ValidationError(f"Paragraph index {paragraph_index} out of range")

        para = self._document.paragraphs[paragraph_index]
        style_name = para.style.name if para.style else ""

        if "List Bullet" in style_name:
            if style_name == "List Bullet 3":
                para.style = "List Bullet 2"
            elif style_name == "List Bullet 2":
                para.style = "List Bullet"
        elif "List Number" in style_name:
            if style_name == "List Number 3":
                para.style = "List Number 2"
            elif style_name == "List Number 2":
                para.style = "List Number"
