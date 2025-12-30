"""Layout handler for page layout operations.

This module provides functionality for managing page layout settings
including margins, page size, headers, footers, and sections.
"""


from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.shared import Inches

from src.core.constants import (
    PAGE_SIZE_A4,
    PAGE_SIZE_LEGAL,
    PAGE_SIZE_LETTER,
)
from src.core.enums import HeaderFooterType, PageOrientation, PageSize, SectionStart
from src.core.exceptions import ValidationError
from src.models.dto import SectionDTO
from src.models.schemas import PageLayout


class LayoutHandler:
    """Handler for page layout operations.

    This class provides methods for managing page layout settings
    in DOCX documents.
    """

    PAGE_SIZE_MAP = {
        PageSize.LETTER: PAGE_SIZE_LETTER,
        PageSize.A4: PAGE_SIZE_A4,
        PageSize.LEGAL: PAGE_SIZE_LEGAL,
    }

    SECTION_START_MAP = {
        SectionStart.CONTINUOUS: WD_SECTION.CONTINUOUS,
        SectionStart.NEW_PAGE: WD_SECTION.NEW_PAGE,
        SectionStart.EVEN_PAGE: WD_SECTION.EVEN_PAGE,
        SectionStart.ODD_PAGE: WD_SECTION.ODD_PAGE,
        SectionStart.NEW_COLUMN: WD_SECTION.NEW_COLUMN,
    }

    def __init__(self, document: Document) -> None:
        """Initialize the layout handler.

        Args:
            document: The Document instance to work with.
        """
        self._document = document

    @property
    def document(self) -> Document:
        """Get the document instance."""
        return self._document

    def get_section(self, index: int) -> SectionDTO:
        """Get a section by index.

        Args:
            index: Section index (0-based).

        Returns:
            Section DTO with layout information.

        Raises:
            ValidationError: If the index is out of range.
        """
        self._validate_section_index(index)
        section = self._document.sections[index]

        return SectionDTO(
            index=index,
            page_width=section.page_width.inches if section.page_width else None,
            page_height=section.page_height.inches if section.page_height else None,
            margin_top=section.top_margin.inches if section.top_margin else None,
            margin_bottom=section.bottom_margin.inches if section.bottom_margin else None,
            margin_left=section.left_margin.inches if section.left_margin else None,
            margin_right=section.right_margin.inches if section.right_margin else None,
            orientation="landscape" if section.orientation == WD_ORIENT.LANDSCAPE else "portrait",
        )

    def get_all_sections(self) -> list[SectionDTO]:
        """Get all sections in the document.

        Returns:
            List of section DTOs.
        """
        return [
            self.get_section(i)
            for i in range(len(self._document.sections))
        ]

    def set_page_layout(
        self,
        section_index: int = 0,
        layout: PageLayout | None = None,
    ) -> SectionDTO:
        """Set page layout for a section.

        Args:
            section_index: Index of the section (default: first section).
            layout: Page layout settings.

        Returns:
            Updated section DTO.

        Raises:
            ValidationError: If the index is out of range.
        """
        self._validate_section_index(section_index)
        section = self._document.sections[section_index]

        if layout:
            # Set page size
            if layout.page_size:
                size = self.PAGE_SIZE_MAP.get(layout.page_size)
                if size:
                    section.page_width = Inches(size[0])
                    section.page_height = Inches(size[1])
            elif layout.width and layout.height:
                section.page_width = Inches(layout.width)
                section.page_height = Inches(layout.height)

            # Set orientation
            if layout.orientation:
                if layout.orientation == PageOrientation.LANDSCAPE:
                    section.orientation = WD_ORIENT.LANDSCAPE
                    # Swap width and height for landscape
                    if section.page_width and section.page_height:
                        if section.page_width < section.page_height:
                            section.page_width, section.page_height = (
                                section.page_height,
                                section.page_width,
                            )
                else:
                    section.orientation = WD_ORIENT.PORTRAIT
                    # Swap back for portrait
                    if section.page_width and section.page_height:
                        if section.page_width > section.page_height:
                            section.page_width, section.page_height = (
                                section.page_height,
                                section.page_width,
                            )

            # Set margins
            if layout.margin_top is not None:
                section.top_margin = Inches(layout.margin_top)
            if layout.margin_bottom is not None:
                section.bottom_margin = Inches(layout.margin_bottom)
            if layout.margin_left is not None:
                section.left_margin = Inches(layout.margin_left)
            if layout.margin_right is not None:
                section.right_margin = Inches(layout.margin_right)

        return self.get_section(section_index)

    def set_margins(
        self,
        section_index: int = 0,
        top: float | None = None,
        bottom: float | None = None,
        left: float | None = None,
        right: float | None = None,
    ) -> SectionDTO:
        """Set page margins for a section.

        Args:
            section_index: Index of the section.
            top: Top margin in inches.
            bottom: Bottom margin in inches.
            left: Left margin in inches.
            right: Right margin in inches.

        Returns:
            Updated section DTO.

        Raises:
            ValidationError: If the index is out of range.
        """
        self._validate_section_index(section_index)
        section = self._document.sections[section_index]

        if top is not None:
            section.top_margin = Inches(top)
        if bottom is not None:
            section.bottom_margin = Inches(bottom)
        if left is not None:
            section.left_margin = Inches(left)
        if right is not None:
            section.right_margin = Inches(right)

        return self.get_section(section_index)

    def add_section(
        self,
        start_type: SectionStart = SectionStart.NEW_PAGE,
    ) -> int:
        """Add a new section to the document.

        Args:
            start_type: Type of section break.

        Returns:
            Index of the new section.
        """
        # Add a paragraph to create a new section
        self._document.add_paragraph()

        # Get the new section and set its start type
        section = self._document.sections[-1]
        section.start_type = self.SECTION_START_MAP.get(
            start_type, WD_SECTION.NEW_PAGE
        )

        return len(self._document.sections) - 1

    def set_header(
        self,
        text: str,
        section_index: int = 0,
        header_type: HeaderFooterType = HeaderFooterType.DEFAULT,
    ) -> None:
        """Set the header text for a section.

        Args:
            text: Header text content.
            section_index: Index of the section.
            header_type: Type of header (default, first, even).

        Raises:
            ValidationError: If the index is out of range.
        """
        self._validate_section_index(section_index)
        section = self._document.sections[section_index]

        if header_type == HeaderFooterType.FIRST:
            section.different_first_page_header_footer = True
            header = section.first_page_header
        elif header_type == HeaderFooterType.EVEN:
            section.different_odd_and_even_page_header_footer = True
            header = section.even_page_header
        else:
            header = section.header

        # Clear existing content and add new text
        header.is_linked_to_previous = False
        for para in header.paragraphs:
            para.clear()
        if header.paragraphs:
            header.paragraphs[0].add_run(text)
        else:
            header.add_paragraph(text)

    def set_footer(
        self,
        text: str,
        section_index: int = 0,
        footer_type: HeaderFooterType = HeaderFooterType.DEFAULT,
    ) -> None:
        """Set the footer text for a section.

        Args:
            text: Footer text content.
            section_index: Index of the section.
            footer_type: Type of footer (default, first, even).

        Raises:
            ValidationError: If the index is out of range.
        """
        self._validate_section_index(section_index)
        section = self._document.sections[section_index]

        if footer_type == HeaderFooterType.FIRST:
            section.different_first_page_header_footer = True
            footer = section.first_page_footer
        elif footer_type == HeaderFooterType.EVEN:
            section.different_odd_and_even_page_header_footer = True
            footer = section.even_page_footer
        else:
            footer = section.footer

        # Clear existing content and add new text
        footer.is_linked_to_previous = False
        for para in footer.paragraphs:
            para.clear()
        if footer.paragraphs:
            footer.paragraphs[0].add_run(text)
        else:
            footer.add_paragraph(text)

    def get_header(
        self,
        section_index: int = 0,
        header_type: HeaderFooterType = HeaderFooterType.DEFAULT,
    ) -> str:
        """Get the header text for a section.

        Args:
            section_index: Index of the section.
            header_type: Type of header.

        Returns:
            Header text content.

        Raises:
            ValidationError: If the index is out of range.
        """
        self._validate_section_index(section_index)
        section = self._document.sections[section_index]

        if header_type == HeaderFooterType.FIRST:
            header = section.first_page_header
        elif header_type == HeaderFooterType.EVEN:
            header = section.even_page_header
        else:
            header = section.header

        return "\n".join(para.text for para in header.paragraphs)

    def get_footer(
        self,
        section_index: int = 0,
        footer_type: HeaderFooterType = HeaderFooterType.DEFAULT,
    ) -> str:
        """Get the footer text for a section.

        Args:
            section_index: Index of the section.
            footer_type: Type of footer.

        Returns:
            Footer text content.

        Raises:
            ValidationError: If the index is out of range.
        """
        self._validate_section_index(section_index)
        section = self._document.sections[section_index]

        if footer_type == HeaderFooterType.FIRST:
            footer = section.first_page_footer
        elif footer_type == HeaderFooterType.EVEN:
            footer = section.even_page_footer
        else:
            footer = section.footer

        return "\n".join(para.text for para in footer.paragraphs)

    def add_page_break(self) -> None:
        """Add a page break at the end of the document."""
        para = self._document.add_paragraph()
        run = para.add_run()
        run.add_break()

    def set_page_numbers(
        self,
        section_index: int = 0,
        start_number: int = 1,
        position: str = "footer",
    ) -> None:
        """Set page numbers for a section.

        Args:
            section_index: Index of the section.
            start_number: Starting page number.
            position: Position ('header' or 'footer').

        Raises:
            ValidationError: If the index is out of range.
        """
        self._validate_section_index(section_index)

        # This would require more complex XML manipulation
        # Placeholder implementation
        text = "Page {PAGE}"
        if position == "header":
            self.set_header(text, section_index)
        else:
            self.set_footer(text, section_index)

    def _validate_section_index(self, index: int) -> None:
        """Validate that a section index is in range.

        Args:
            index: Section index to validate.

        Raises:
            ValidationError: If the index is out of range.
        """
        if index < 0 or index >= len(self._document.sections):
            raise ValidationError(
                f"Section index {index} out of range (0-{len(self._document.sections) - 1})"
            )
