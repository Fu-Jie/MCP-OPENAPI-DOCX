"""Pytest configuration and fixtures.

This module provides shared fixtures and configuration
for the test suite.
"""

import os
import sys
import pytest
import asyncio
from typing import AsyncGenerator, Generator
from pathlib import Path

# Add project root to path
project_root = Path(__file__).parent.parent
sys.path.insert(0, str(project_root))

from sqlalchemy.ext.asyncio import AsyncSession, create_async_engine, async_sessionmaker
from httpx import AsyncClient

from src.core.config import Settings, get_settings
from src.database.base import Base
from src.api.main import create_application


# Test database URL (SQLite for testing)
TEST_DATABASE_URL = "sqlite+aiosqlite:///./test_db.sqlite"


@pytest.fixture(scope="session")
def event_loop() -> Generator:
    """Create event loop for async tests."""
    loop = asyncio.get_event_loop_policy().new_event_loop()
    yield loop
    loop.close()


@pytest.fixture(scope="session")
def test_settings() -> Settings:
    """Get test settings."""
    return Settings(
        database_url=TEST_DATABASE_URL,
        debug=True,
        environment="test",
        secret_key="test-secret-key-for-testing-only",
        upload_dir="./test_uploads",
        export_dir="./test_exports",
        temp_dir="./test_temp",
    )


@pytest.fixture(scope="session")
async def test_engine(test_settings: Settings):
    """Create test database engine."""
    engine = create_async_engine(
        test_settings.database_url,
        echo=False,
    )

    async with engine.begin() as conn:
        await conn.run_sync(Base.metadata.create_all)

    yield engine

    async with engine.begin() as conn:
        await conn.run_sync(Base.metadata.drop_all)

    await engine.dispose()


@pytest.fixture
async def db_session(test_engine) -> AsyncGenerator[AsyncSession, None]:
    """Get database session for tests."""
    async_session = async_sessionmaker(
        test_engine,
        class_=AsyncSession,
        expire_on_commit=False,
    )

    async with async_session() as session:
        yield session
        await session.rollback()


@pytest.fixture
def app():
    """Create test application."""
    return create_application()


@pytest.fixture
async def client(app) -> AsyncGenerator[AsyncClient, None]:
    """Get async test client."""
    async with AsyncClient(app=app, base_url="http://test") as ac:
        yield ac


@pytest.fixture(scope="session", autouse=True)
def create_test_directories(test_settings: Settings):
    """Create test directories."""
    for dir_path in [
        test_settings.upload_dir,
        test_settings.export_dir,
        test_settings.temp_dir,
    ]:
        os.makedirs(dir_path, exist_ok=True)

    yield

    # Cleanup
    import shutil
    for dir_path in [
        test_settings.upload_dir,
        test_settings.export_dir,
        test_settings.temp_dir,
    ]:
        if os.path.exists(dir_path):
            shutil.rmtree(dir_path)


@pytest.fixture
def sample_docx_content() -> bytes:
    """Get sample DOCX file content."""
    from docx import Document
    from io import BytesIO

    doc = Document()
    doc.add_heading("Test Document", 0)
    doc.add_paragraph("This is a test paragraph.")
    doc.add_paragraph("This is another paragraph.")

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


@pytest.fixture
def sample_document_path(sample_docx_content, test_settings) -> str:
    """Create a sample document file for testing."""
    file_path = os.path.join(test_settings.upload_dir, "test_doc.docx")
    with open(file_path, "wb") as f:
        f.write(sample_docx_content)
    return file_path


@pytest.fixture
def mock_user() -> dict:
    """Get mock user data."""
    return {
        "id": "test-user-id-123",
        "email": "test@example.com",
        "username": "testuser",
        "full_name": "Test User",
        "is_active": True,
        "is_superuser": False,
    }


# Marker for slow tests
def pytest_configure(config):
    """Configure pytest markers."""
    config.addinivalue_line("markers", "slow: marks tests as slow")
    config.addinivalue_line("markers", "integration: marks tests as integration tests")
