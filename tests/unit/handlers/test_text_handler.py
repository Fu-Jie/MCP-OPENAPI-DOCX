"""Unit tests for text handler."""

import pytest
from docx import Document

from src.handlers.text_handler import TextHandler


class TestTextHandler:
    """Test cases for TextHandler class."""

    def setup_method(self):
        """Set up test fixtures."""
        self.handler = TextHandler()
        self.doc = Document()

    def test_add_paragraph(self):
        """Test adding a paragraph."""
        para = self.handler.add_paragraph(self.doc, "Test paragraph")
        assert para is not None
        assert para.text == "Test paragraph"

    def test_add_paragraph_with_style(self):
        """Test adding a paragraph with style."""
        para = self.handler.add_paragraph(
            self.doc, "Heading text", style="Heading 1"
        )
        assert para is not None
        assert para.text == "Heading text"

    def test_get_paragraphs(self):
        """Test getting paragraphs."""
        self.handler.add_paragraph(self.doc, "Para 1")
        self.handler.add_paragraph(self.doc, "Para 2")

        paragraphs = self.handler.get_paragraphs(self.doc)
        assert len(paragraphs) == 2
        assert paragraphs[0]["text"] == "Para 1"
        assert paragraphs[1]["text"] == "Para 2"

    def test_update_paragraph(self):
        """Test updating a paragraph."""
        self.handler.add_paragraph(self.doc, "Original text")
        self.handler.update_paragraph(self.doc, 0, text="Updated text")

        paragraphs = self.handler.get_paragraphs(self.doc)
        assert paragraphs[0]["text"] == "Updated text"

    def test_delete_paragraph(self):
        """Test deleting a paragraph."""
        self.handler.add_paragraph(self.doc, "Para 1")
        self.handler.add_paragraph(self.doc, "Para 2")
        self.handler.add_paragraph(self.doc, "Para 3")

        self.handler.delete_paragraph(self.doc, 1)

        paragraphs = self.handler.get_paragraphs(self.doc)
        assert len(paragraphs) == 2

    def test_format_text_bold(self):
        """Test formatting text as bold."""
        para = self.handler.add_paragraph(self.doc, "Bold text")
        self.handler.format_text(self.doc, 0, bold=True)

        # Check that formatting was applied
        assert para.runs[0].bold is True

    def test_format_text_italic(self):
        """Test formatting text as italic."""
        para = self.handler.add_paragraph(self.doc, "Italic text")
        self.handler.format_text(self.doc, 0, italic=True)

        assert para.runs[0].italic is True

    def test_search_text(self):
        """Test searching for text."""
        self.handler.add_paragraph(self.doc, "Hello world")
        self.handler.add_paragraph(self.doc, "World is beautiful")
        self.handler.add_paragraph(self.doc, "Hello again")

        results = self.handler.search_text(self.doc, "world")
        assert len(results) == 2

    def test_search_text_case_sensitive(self):
        """Test case-sensitive text search."""
        self.handler.add_paragraph(self.doc, "Hello World")
        self.handler.add_paragraph(self.doc, "hello world")

        results = self.handler.search_text(self.doc, "World", case_sensitive=True)
        assert len(results) == 1

    def test_replace_text(self):
        """Test replacing text."""
        self.handler.add_paragraph(self.doc, "Hello world")
        self.handler.add_paragraph(self.doc, "world is great")

        count = self.handler.replace_text(self.doc, "world", "universe")
        assert count == 2

        paragraphs = self.handler.get_paragraphs(self.doc)
        assert "universe" in paragraphs[0]["text"]
        assert "universe" in paragraphs[1]["text"]
